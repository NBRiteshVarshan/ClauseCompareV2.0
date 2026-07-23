import streamlit as st
import plotly.graph_objects as go
from datetime import datetime
import html
import time
import re
import io
from docx import Document
from docx.shared import Inches

from document_processor import ClauseExtractor, get_document_summary
from clause_matcher import LegalClauseMatcher
from utils import format_report, save_report, generate_pdf_report, categorize_results
from db_manager import DatabaseManager

st.set_page_config(
    page_title="ClauseCompare V2.0",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Enhanced Custom CSS
st.markdown("""
    <style>
    .main-header {
        font-size: 2.8rem;
        font-weight: 700;
        color: #1E3A8A;
        text-align: center;
        margin-bottom: 1rem;
        letter-spacing: -0.5px;
    }
    .sub-header {
        color: #6B7280;
        text-align: center;
        margin-bottom: 2rem;
        font-size: 1.1rem;
    }
    .stats-grid {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 1rem;
        margin-bottom: 2rem;
    }
    .stat-card {
        background: white;
        padding: 1.5rem 1rem;
        border-radius: 1rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
        text-align: center;
        border: 1px solid #F3F4F6;
        transition: transform 0.2s;
    }
    .stat-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 20px rgba(0,0,0,0.08);
    }
    .stat-number {
        font-size: 2.8rem;
        font-weight: 700;
        line-height: 1.2;
    }
    .stat-label {
        color: #6B7280;
        font-size: 0.9rem;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    .doc-card {
        background: white;
        padding: 1rem 1.25rem;
        border-radius: 0.75rem;
        margin: 0.5rem 0;
        border: 1px solid #E5E7EB;
        display: flex;
        justify-content: space-between;
        align-items: center;
        transition: border-color 0.2s;
    }
    .doc-card:hover {
        border-color: #3B82F6;
    }
    .doc-name {
        font-weight: 500;
        color: #111827;
    }
    .doc-meta {
        color: #6B7280;
        font-size: 0.85rem;
    }
    .status-badge {
        padding: 0.15rem 0.6rem;
        border-radius: 9999px;
        font-size: 0.75rem;
        font-weight: 500;
    }
    .status-compared {
        background: #D1FAE5;
        color: #065F46;
    }
    .status-pending {
        background: #FEF3C7;
        color: #92400E;
    }
    .delete-btn {
        background: none;
        border: none;
        color: #9CA3AF;
        font-size: 1.1rem;
        cursor: pointer;
        padding: 0 0.2rem;
        transition: color 0.2s;
    }
    .delete-btn:hover {
        color: #EF4444;
    }
    .exact-box {
        background: #F0FDF4;
        padding: 1rem;
        border-radius: 0.75rem;
        margin: 0.5rem 0;
        border-left: 4px solid #059669;
    }
    .partial-box {
        background: #FEF3C7;
        padding: 1rem;
        border-radius: 0.75rem;
        margin: 0.5rem 0;
        border-left: 4px solid #D97706;
    }
    .unique-box {
        background: #FEF2F2;
        padding: 1rem;
        border-radius: 0.75rem;
        margin: 0.5rem 0;
        border-left: 4px solid #DC2626;
    }
    .stButton > button {
        background-color: #1E3A8A;
        color: white;
        font-weight: 500;
        border: none;
        border-radius: 0.5rem;
        padding: 0.5rem 1.5rem;
        transition: background 0.2s;
    }
    .stButton > button:hover {
        background-color: #1E40AF;
        color: white;
    }
    .stButton > button:disabled {
        background-color: #9CA3AF;
    }
    .sidebar-title {
        font-size: 1.5rem;
        font-weight: 700;
        color: #1E3A8A;
        margin-bottom: 0.25rem;
    }
    .sidebar-sub {
        color: #6B7280;
        font-size: 0.9rem;
        margin-bottom: 1rem;
    }
    .footer {
        text-align: center;
        color: #9CA3AF;
        font-size: 0.8rem;
        margin-top: 3rem;
        border-top: 1px solid #E5E7EB;
        padding-top: 1.5rem;
    }
    .variable-input {
        margin-bottom: 0.75rem;
    }
    @media (max-width: 640px) {
        .stats-grid {
            grid-template-columns: 1fr;
        }
    }
    </style>
""", unsafe_allow_html=True)

def init_session_state():
    if 'db' not in st.session_state:
        st.session_state.db = DatabaseManager()
    if 'page' not in st.session_state:
        st.session_state.page = 'Dashboard'
    if 'comparison_results' not in st.session_state:
        st.session_state.comparison_results = None
    if 'compare_doc1_id' not in st.session_state:
        st.session_state.compare_doc1_id = None
    if 'compare_doc2_id' not in st.session_state:
        st.session_state.compare_doc2_id = None
    if 'compare_doc1_clauses' not in st.session_state:
        st.session_state.compare_doc1_clauses = None
    if 'compare_doc2_clauses' not in st.session_state:
        st.session_state.compare_doc2_clauses = None
    if 'confirm_delete_id' not in st.session_state:
        st.session_state.confirm_delete_id = None
    # Document Generator state
    if 'template_variables' not in st.session_state:
        st.session_state.template_variables = []
    if 'template_file' not in st.session_state:
        st.session_state.template_file = None
    if 'generated_docx' not in st.session_state:
        st.session_state.generated_docx = None

def escape_text(text: str) -> str:
    return html.escape(text)

def render_stats(total, compared, pending):
    st.markdown(f"""
        <div class="stats-grid">
            <div class="stat-card">
                <div class="stat-number" style="color:#4F46E5;">{total}</div>
                <div class="stat-label">📄 Total Documents</div>
            </div>
            <div class="stat-card">
                <div class="stat-number" style="color:#059669;">{compared}</div>
                <div class="stat-label">✅ Compared</div>
            </div>
            <div class="stat-card">
                <div class="stat-number" style="color:#D97706;">{pending}</div>
                <div class="stat-label">⏳ Pending</div>
            </div>
        </div>
    """, unsafe_allow_html=True)

def render_doc_list(docs, title, db):
    if not docs:
        st.info(f"No {title.lower()} documents uploaded yet.")
        return
    for doc in docs:
        status = "Compared" if doc['is_compared'] else "Pending"
        status_class = "status-compared" if doc['is_compared'] else "status-pending"
        with st.container():
            col1, col2 = st.columns([10, 1])
            with col1:
                st.markdown(f"""
                    <div class="doc-card">
                        <div>
                            <span class="doc-name">{doc['name']}</span>
                            <span class="status-badge {status_class}">{status}</span>
                            <span class="doc-meta"> • Uploaded: {doc['uploaded_at']}</span>
                        </div>
                    </div>
                """, unsafe_allow_html=True)
            with col2:
                if st.button("🗑️", key=f"del_{doc['id']}"):
                    st.session_state.confirm_delete_id = doc['id']
                    st.rerun()

def extract_variables(text: str) -> list:
    """Extract all unique variable names from $#...#$ pattern."""
    pattern = r'\$#(.*?)#\$'
    matches = re.findall(pattern, text)
    # Remove duplicates while preserving order
    seen = set()
    unique = []
    for m in matches:
        if m not in seen:
            seen.add(m)
            unique.append(m)
    return unique

def process_docx_template(file_bytes: bytes) -> tuple:
    """Read docx, extract variables and return (document, variables, text)."""
    doc = Document(io.BytesIO(file_bytes))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    text = "\n".join(full_text)
    variables = extract_variables(text)
    return doc, variables, text

def generate_docx_from_template(template_bytes: bytes, replacements: dict) -> bytes:
    """Replace $#...#$ in the docx with values and return as bytes."""
    doc = Document(io.BytesIO(template_bytes))
    # Replace in paragraphs
    for para in doc.paragraphs:
        for key, value in replacements.items():
            # We need to replace the exact pattern with value
            pattern = f"$#{key}#$"
            if pattern in para.text:
                # Replace inline runs carefully
                # For simplicity, we can replace the entire paragraph text
                # but that would lose formatting. Better to replace run text.
                # We'll iterate through runs
                for run in para.runs:
                    if pattern in run.text:
                        run.text = run.text.replace(pattern, value)
                # Also handle paragraphs that might have the pattern split across runs
                # For safety, also replace at paragraph level if runs didn't catch all
                if pattern in para.text:
                    para.text = para.text.replace(pattern, value)
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, value in replacements.items():
                            pattern = f"$#{key}#$"
                            if pattern in run.text:
                                run.text = run.text.replace(pattern, value)
                    # Fallback paragraph level
                    for key, value in replacements.items():
                        pattern = f"$#{key}#$"
                        if pattern in para.text:
                            para.text = para.text.replace(pattern, value)
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def main():
    init_session_state()
    db = st.session_state.db

    with st.sidebar:
        st.markdown('<div class="sidebar-title">📂 ClauseCompare</div>', unsafe_allow_html=True)
        st.markdown('<div class="sidebar-sub">Your local contract companion</div>', unsafe_allow_html=True)
        st.divider()
        page = st.radio("Navigate", ["Dashboard", "Compare Documents", "Add Document", "Document Generator", "Reports"])
        st.session_state.page = page
        st.divider()
        if st.button("🔄 Reset App", use_container_width=True):
            st.cache_data.clear()
            st.cache_resource.clear()
            for key in list(st.session_state.keys()):
                if key != 'db':
                    del st.session_state[key]
            st.rerun()
        st.divider()
        st.markdown("""
            <div style="font-size:0.8rem; color:#6B7280; text-align:center;">
                🔒 All data stays local<br>
                ⚡ Powered by local LLM
            </div>
        """, unsafe_allow_html=True)

    # ---------- DASHBOARD ----------
    if st.session_state.page == "Dashboard":
        st.markdown('<h1 class="main-header">📊 Dashboard</h1>', unsafe_allow_html=True)
        st.markdown('<div class="sub-header">Overview of your contract repository</div>', unsafe_allow_html=True)

        stats = db.get_stats()
        render_stats(stats['total'], stats['compared'], stats['pending'])

        # Delete confirmation
        if st.session_state.confirm_delete_id:
            with st.expander("⚠️ Confirm Delete", expanded=True):
                doc = db.get_document(st.session_state.confirm_delete_id)
                if doc:
                    st.warning(f"Are you sure you want to delete **{doc['name']}**? This action cannot be undone.")
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("✅ Yes, Delete", use_container_width=True):
                            db.delete_document(st.session_state.confirm_delete_id)
                            st.session_state.confirm_delete_id = None
                            st.success("Document deleted.")
                            st.rerun()
                    with col2:
                        if st.button("❌ Cancel", use_container_width=True):
                            st.session_state.confirm_delete_id = None
                            st.rerun()
                else:
                    st.session_state.confirm_delete_id = None
                    st.rerun()

        st.divider()
        st.subheader("📄 Company Documents")
        company_docs = db.get_all_documents('company')
        render_doc_list(company_docs, "company", db)

        st.subheader("📄 Third‑Party Documents")
        third_party_docs = db.get_all_documents('third_party')
        render_doc_list(third_party_docs, "third‑party", db)

    # ---------- ADD DOCUMENT ----------
    elif st.session_state.page == "Add Document":
        st.markdown('<h1 class="main-header">📤 Add New Document</h1>', unsafe_allow_html=True)
        st.markdown('<div class="sub-header">Upload a PDF or DOCX to extract clauses for comparison</div>', unsafe_allow_html=True)

        with st.form("upload_form", clear_on_submit=False):
            col1, col2 = st.columns([3, 1])
            with col1:
                doc_name = st.text_input("Document Name (optional)", placeholder="e.g., NDA_v3")
            with col2:
                category = st.selectbox("Category", ["company", "third_party"])
            uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'docx'])
            submitted = st.form_submit_button("📤 Upload & Extract", use_container_width=True)

            if submitted and uploaded_file is not None:
                with st.spinner("Extracting clauses..."):
                    extractor = ClauseExtractor()
                    try:
                        file_bytes = uploaded_file.read()
                        if uploaded_file.type == 'application/pdf':
                            clauses = extractor.extract_from_pdf(file_bytes)
                        else:
                            clauses = extractor.extract_from_docx(file_bytes)
                        doc_id = db.add_document(doc_name or uploaded_file.name, category)
                        db.add_clauses(doc_id, clauses)
                        st.success(f"✅ **{uploaded_file.name}** uploaded successfully with **{len(clauses)}** clauses!")
                        st.info("Return to the Dashboard to see it listed.")
                    except Exception as e:
                        st.error(f"❌ Error: {e}")

    # ---------- COMPARE DOCUMENTS ----------
    elif st.session_state.page == "Compare Documents":
        st.markdown('<h1 class="main-header">⚖️ Compare Documents</h1>', unsafe_allow_html=True)
        st.markdown('<div class="sub-header">Select two documents to run a semantic comparison</div>', unsafe_allow_html=True)

        company_docs = db.get_all_documents('company')
        third_party_docs = db.get_all_documents('third_party')

        if len(company_docs) < 1 or len(third_party_docs) < 1:
            st.warning("You need at least **one Company document** and **one Third‑Party document** to compare. Upload them first.")
        else:
            comp_options = {f"{d['name']} (ID: {d['id']})": d['id'] for d in company_docs}
            third_options = {f"{d['name']} (ID: {d['id']})": d['id'] for d in third_party_docs}

            col1, col2 = st.columns(2)
            with col1:
                selected1 = st.selectbox("📄 Select Company Document", options=list(comp_options.keys()))
                doc1_id = comp_options[selected1]
            with col2:
                selected2 = st.selectbox("📄 Select Third‑Party Document", options=list(third_options.keys()))
                doc2_id = third_options[selected2]

            comparison_name = st.text_input(
                "✏️ Name this comparison (optional)",
                placeholder="e.g., NDA v3 vs Vendor A",
                help="Give it a meaningful name for easy reference in the Reports section."
            )

            if st.button("🚀 Run Comparison", use_container_width=True):
                with st.spinner("Loading clauses and comparing..."):
                    clauses1 = db.get_clauses(doc1_id)
                    clauses2 = db.get_clauses(doc2_id)
                    if not clauses1 or not clauses2:
                        st.error("One of the documents has no clauses. Please re‑upload.")
                    else:
                        st.session_state.compare_doc1_id = doc1_id
                        st.session_state.compare_doc2_id = doc2_id
                        st.session_state.compare_doc1_clauses = clauses1
                        st.session_state.compare_doc2_clauses = clauses2

                        doc1 = [{'number': c['clause_number'], 'text': c['text'], 'metadata': c['metadata']} for c in clauses1]
                        doc2 = [{'number': c['clause_number'], 'text': c['text'], 'metadata': c['metadata']} for c in clauses2]

                        matcher = LegalClauseMatcher()
                        progress_placeholder = st.empty()
                        status_placeholder = st.empty()
                        def update_progress(prog, status):
                            progress_placeholder.progress(prog)
                            status_placeholder.text(status)

                        results = matcher.match_documents(
                            doc1, doc2,
                            similarity_threshold=0.4,
                            high_similarity_threshold=0.85,
                            match_threshold=0.5,
                            progress_callback=update_progress
                        )
                        result_id = db.save_comparison_result(doc1_id, doc2_id, results, comparison_name)
                        db.update_document_compared(doc1_id, result_id)
                        db.update_document_compared(doc2_id, result_id)
                        st.session_state.comparison_results = results
                        st.success("✅ Comparison complete! See results below.")
                        st.rerun()

            # Display results if they exist
            if st.session_state.comparison_results is not None:
                results = st.session_state.comparison_results
                clauses1 = st.session_state.compare_doc1_clauses
                clauses2 = st.session_state.compare_doc2_clauses
                if clauses1 is not None and clauses2 is not None:
                    doc1 = [{'number': c['clause_number'], 'text': c['text'], 'metadata': c['metadata']} for c in clauses1]
                    doc2 = [{'number': c['clause_number'], 'text': c['text'], 'metadata': c['metadata']} for c in clauses2]
                else:
                    doc1_id = st.session_state.compare_doc1_id
                    doc2_id = st.session_state.compare_doc2_id
                    if doc1_id and doc2_id:
                        clauses1 = db.get_clauses(doc1_id)
                        clauses2 = db.get_clauses(doc2_id)
                        if clauses1 and clauses2:
                            doc1 = [{'number': c['clause_number'], 'text': c['text'], 'metadata': c['metadata']} for c in clauses1]
                            doc2 = [{'number': c['clause_number'], 'text': c['text'], 'metadata': c['metadata']} for c in clauses2]
                        else:
                            doc1 = doc2 = []
                    else:
                        doc1 = doc2 = []

                if doc1 and doc2:
                    st.subheader("📊 Comparison Results")

                    exact, partial, unique = categorize_results(results, doc1, doc2)
                    total_matched = len(exact) + len(partial)
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Total Clauses", len(doc1) + len(doc2))
                    col2.metric("Matched", total_matched)
                    col3.metric("Unique", len(unique))

                    with st.expander(f"✅ Exact Matches (Similarity ≥ 0.999) — {len(exact)} pairs"):
                        if exact:
                            for m in exact:
                                safe_doc1 = escape_text(m['doc1_text'])
                                safe_doc2 = escape_text(m['doc2_text'])
                                st.markdown(f"""
                                    <div class="exact-box">
                                        <strong>📄 Doc1 – Clause {m['doc1_num']}</strong><br>{safe_doc1}<br><br>
                                        <strong>📄 Doc2 – Clause {m['doc2_num']}</strong><br>{safe_doc2}<br><br>
                                        <span style="color:#059669;">✅ Similarity: {m['similarity']:.3f}</span>
                                    </div>
                                """, unsafe_allow_html=True)
                        else:
                            st.info("No exact matches found.")

                    with st.expander(f"🟡 Partial Matches (0.5 ≤ Similarity < 0.999) — {len(partial)} pairs"):
                        if partial:
                            for m in partial:
                                safe_doc1 = escape_text(m['doc1_text'])
                                safe_doc2 = escape_text(m['doc2_text'])
                                st.markdown(f"""
                                    <div class="partial-box">
                                        <strong>📄 Doc1 – Clause {m['doc1_num']}</strong><br>{safe_doc1}<br><br>
                                        <strong>📄 Doc2 – Clause {m['doc2_num']}</strong><br>{safe_doc2}<br><br>
                                        <span style="color:#D97706;">🔶 Similarity: {m['similarity']:.3f}</span>
                                    </div>
                                """, unsafe_allow_html=True)
                        else:
                            st.info("No partial matches found.")

                    with st.expander(f"🔴 Unique Clauses (not matched) — {len(unique)} clauses"):
                        if unique:
                            unique_doc1_list = [u for u in unique if u['document'] == 'Document 1']
                            unique_doc2_list = [u for u in unique if u['document'] == 'Document 2']
                            if unique_doc1_list:
                                st.markdown("**Document 1 Unique:**")
                                for u in unique_doc1_list:
                                    safe_text = escape_text(u['text'])
                                    st.markdown(f"""
                                        <div class="unique-box">
                                            <strong>Clause {u['number']}</strong><br>{safe_text}<br>
                                            <span style="color:#DC2626;">❌ Best similarity: {u['similarity']:.3f}</span>
                                        </div>
                                    """, unsafe_allow_html=True)
                            if unique_doc2_list:
                                st.markdown("**Document 2 Unique:**")
                                for u in unique_doc2_list:
                                    safe_text = escape_text(u['text'])
                                    st.markdown(f"""
                                        <div class="unique-box">
                                            <strong>Clause {u['number']}</strong><br>{safe_text}<br>
                                            <span style="color:#DC2626;">❌ Best similarity: {u['similarity']:.3f}</span>
                                        </div>
                                    """, unsafe_allow_html=True)
                        else:
                            st.info("All clauses are matched (no unique clauses).")

                    st.subheader("📥 Download Reports")
                    txt_report = format_report(results)
                    json_report = save_report(results)
                    pdf_report = generate_pdf_report(results, doc1, doc2)
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.download_button("📄 PDF", pdf_report, file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf", mime="application/pdf")
                    with col2:
                        st.download_button("📄 TXT", txt_report, file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
                    with col3:
                        st.download_button("📊 JSON", json_report, file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")

    # ---------- DOCUMENT GENERATOR ----------
    elif st.session_state.page == "Document Generator":
        st.markdown('<h1 class="main-header">📄 Document Generator</h1>', unsafe_allow_html=True)
        st.markdown('<div class="sub-header">Upload a DOCX template with variables in $#variable_name#$ format</div>', unsafe_allow_html=True)

        uploaded_template = st.file_uploader("Upload a DOCX template", type=['docx'])

        if uploaded_template is not None:
            # Read the file
            file_bytes = uploaded_template.read()
            doc, variables, full_text = process_docx_template(file_bytes)

            # Store in session state for later use
            st.session_state.template_bytes = file_bytes
            st.session_state.template_variables = variables

            if not variables:
                st.info("No variables found in the template. The document contains no $#...#$ patterns.")
                # Offer to download the original as is
                st.download_button(
                    label="📥 Download Template (as is)",
                    data=file_bytes,
                    file_name=f"template_{uploaded_template.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.success(f"Found **{len(variables)}** unique variable(s): {', '.join(variables)}")

                # Form for variable values
                with st.form("variable_form"):
                    st.markdown("### Enter values for each variable")
                    values = {}
                    for var in variables:
                        values[var] = st.text_input(f"${var}", key=f"var_{var}")
                    submitted = st.form_submit_button("📝 Generate Document")

                # Process submission outside the form
                if submitted:
                    missing = [var for var, val in values.items() if not val.strip()]
                    if missing:
                        st.error(f"Please fill in values for: {', '.join(missing)}")
                    else:
                        replacements = {var: values[var].strip() for var in variables}
                        generated_bytes = generate_docx_from_template(file_bytes, replacements)
                        st.session_state.generated_docx = generated_bytes
                        st.success("✅ Document generated successfully!")

                # Display download button if generated bytes exist (outside form)
                if st.session_state.get('generated_docx'):
                    st.download_button(
                        label="📥 Download Generated Document",
                        data=st.session_state.generated_docx,
                        file_name=f"generated_{uploaded_template.name}",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                # Option to clear the generated doc
                if st.button("🔄 Clear generated document"):
                    st.session_state.generated_docx = None
                    st.rerun()

        else:
            st.info("Please upload a DOCX template to get started.")
            st.markdown("""
                **How it works:**
                1. Create a DOCX file with placeholders like `$#ClientName#$` and `$#EffectiveDate#$`.
                2. Upload it here.
                3. Enter the values for each variable.
                4. Download the populated document with all formatting preserved.
            """)
    
    # ---------- REPORTS ----------
    elif st.session_state.page == "Reports":
        st.markdown('<h1 class="main-header">📜 Past Comparisons</h1>', unsafe_allow_html=True)
        st.markdown('<div class="sub-header">All previously saved comparisons</div>', unsafe_allow_html=True)

        comparisons = db.get_all_comparisons()
        if not comparisons:
            st.info("No past comparisons yet. Run your first comparison to see results here.")
        else:
            for comp in comparisons:
                display_name = comp['comparison_name'] if comp['comparison_name'] else f"Comparison #{comp['id']}"
                with st.expander(f"{display_name} – {comp['doc1_name']} vs {comp['doc2_name']} ({comp['timestamp']})"):
                    full = db.get_comparison_with_docs(comp['id'])
                    if full:
                        result_json = full['result_json']
                        clauses1 = db.get_clauses(full['doc1_id'])
                        clauses2 = db.get_clauses(full['doc2_id'])
                        doc1 = [{'number': c['clause_number'], 'text': c['text'], 'metadata': c['metadata']} for c in clauses1]
                        doc2 = [{'number': c['clause_number'], 'text': c['text'], 'metadata': c['metadata']} for c in clauses2]
                        if doc1 and doc2:
                            txt_report = format_report(result_json)
                            json_report = save_report(result_json)
                            pdf_report = generate_pdf_report(result_json, doc1, doc2)
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.download_button("📄 PDF", pdf_report, file_name=f"comparison_{comp['id']}.pdf", mime="application/pdf")
                            with col2:
                                st.download_button("📄 TXT", txt_report, file_name=f"comparison_{comp['id']}.txt")
                            with col3:
                                st.download_button("📊 JSON", json_report, file_name=f"comparison_{comp['id']}.json")
                        else:
                            st.warning("Could not retrieve clauses for this comparison.")
                    else:
                        st.warning("Could not load comparison data.")

    # Footer
    st.divider()
    st.markdown("""
        <div class="footer">
            ClauseCompare V2.0 • All processing runs locally • No data leaves your machine
        </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    try:
        import ollama
        ollama.list()
    except Exception:
        st.warning("⚠️ Ollama not detected. Please start Ollama and pull llama3.2:3b.")
        st.info("""
        To install:
        1. Download Ollama from https://ollama.ai
        2. Run in terminal: `ollama pull llama3.2:3b`
        3. Run in terminal: `ollama serve`
        """)
    main()